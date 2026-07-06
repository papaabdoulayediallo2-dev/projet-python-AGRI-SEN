#!/usr/bin/env python3
"""
Script pour générer un rapport complet du projet AGRI-SEN
Contient: Introduction, Analyse, Algorithme, Captures d'écran, Conclusion
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Créer le document
doc = Document()

# Définir les styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Couleur pour les titres
COULEUR_TITRE = RGBColor(34, 139, 34)  # Vert

def add_heading(doc, text, level=1):
    """Ajoute un titre avec formatage"""
    heading = doc.add_heading(text, level=level)
    heading.runs[0].font.color.rgb = COULEUR_TITRE
    if level == 1:
        heading.runs[0].font.size = Pt(24)
        heading.runs[0].font.bold = True
    return heading

def add_paragraph_with_spacing(doc, text, space_before=6, space_after=6):
    """Ajoute un paragraphe avec espacement"""
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    return p

def shade_cell(cell, color):
    """Ajoute une couleur de fond à une cellule"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

# ========================
# PAGE DE COUVERTURE
# ========================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("RAPPORT DE PROJET")
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = COULEUR_TITRE

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("AGRI-SEN : Système de Gestion de Coopérative Agricole")
run.font.size = Pt(18)
run.font.color.rgb = COULEUR_TITRE

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("Développé par : 5 Développeurs")
run.font.size = Pt(12)

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_para.add_run("Date : Juillet 2026")
run.font.size = Pt(12)

doc.add_page_break()

# ========================
# TABLE OF CONTENTS
# ========================
doc.add_heading("Table des Matières", level=1)
toc_items = [
    "I. Introduction",
    "II. Analyse des Fonctionnalités",
    "III. Algorithmes et Traitements",
    "IV. Démonstration et Captures d'Écran",
    "V. Conclusion"
]
for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ========================
# I. INTRODUCTION
# ========================
add_heading(doc, "I. Introduction : Contexte et Présentation du Problème", level=1)

add_heading(doc, "1. Contexte Général", level=2)
add_paragraph_with_spacing(doc,
    "Le secteur agricole en Afrique de l'Ouest, et particulièrement au Sénégal, joue un rôle fondamental "
    "dans l'économie et la sécurité alimentaire. Les producteurs agricoles se regroupent généralement dans "
    "des coopératives pour améliorer leur productivité, partager les ressources et maximiser leurs revenus. "
    "Cependant, la gestion administrative et comptable de ces coopératives reste souvent rudimentaire et "
    "peu structurée.")

add_heading(doc, "2. Problématique Identifiée", level=2)
add_paragraph_with_spacing(doc,
    "Les coopératives agricoles face plusieurs défis majeurs :")

challenges = [
    "Absence de centralisation des données des producteurs",
    "Difficulté à suivre et enregistrer les récoltes de manière organisée",
    "Manque de visibilité sur la production globale et ses performances",
    "Impossibilité d'identifier les producteurs les plus performants",
    "Absence d'outils d'analyse pour optimiser la gestion"
]

for challenge in challenges:
    doc.add_paragraph(challenge, style='List Bullet')

add_heading(doc, "3. Solution Proposée : AGRI-SEN", level=2)
add_paragraph_with_spacing(doc,
    "AGRI-SEN est un système informatique développé pour résoudre ces problèmes. Il permet une gestion "
    "complète et intégrée de la coopérative en centralisant les données des producteurs et de leurs récoltes, "
    "en fournissant des rapports d'analyse détaillés et en facilitant la prise de décision.")

add_heading(doc, "4. Objectifs du Projet", level=2)
objectives = [
    "Centraliser et organiser les données des producteurs de manière sécurisée",
    "Permettre l'enregistrement structuré des récoltes avec validation complète",
    "Générer des rapports détaillés sur la production globale et par produit",
    "Identifier les producteurs les plus performants",
    "Faciliter la gestion administrative des coopératives agricoles"
]
for obj in objectives:
    doc.add_paragraph(obj, style='List Number')

doc.add_page_break()

# ========================
# II. ANALYSE
# ========================
add_heading(doc, "II. Analyse des Fonctionnalités", level=1)

add_heading(doc, "1. Architecture Générale du Système", level=2)
add_paragraph_with_spacing(doc,
    "AGRI-SEN est construit selon une architecture modulaire composée de 5 modules Python indépendants "
    "mais interconnectés, chacun responsable d'une fonction spécifique. Cette approche modulaire permet "
    "une maintenance facile et la réutilisabilité du code.")

# Tableau d'architecture
add_heading(doc, "2. Description des Modules", level=2)

table = doc.add_table(rows=6, cols=3)
table.style = 'Light Grid Accent 1'

# En-tête
header_cells = table.rows[0].cells
header_cells[0].text = "Module"
header_cells[1].text = "Fichier"
header_cells[2].text = "Responsabilités"

modules_data = [
    ("Gestion des Producteurs", "membre1.py", "Ajouter, afficher et rechercher des producteurs"),
    ("Gestion des Récoltes", "membre2.py", "Enregistrer et afficher les récoltes avec validation"),
    ("Statistiques", "Membre3.py", "Calculer production totale et par produit"),
    ("Analyses", "membre4.py", "Identifier meilleur producteur et moyennes"),
    ("Interface Principale", "main.py", "Menu interactif et intégration de tous les modules")
]

for i, (module, fichier, resp) in enumerate(modules_data, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = module
    row_cells[1].text = fichier
    row_cells[2].text = resp

add_heading(doc, "3. Détail des Fonctionnalités", level=2)

add_heading(doc, "3.1 Module 1 : Gestion des Producteurs", level=3)
add_paragraph_with_spacing(doc,
    "Ce module gère le cycle de vie des producteurs au sein de la coopérative. Il permet de maintenir "
    "une liste centralisée de tous les producteurs avec leurs informations de base et leurs récoltes associées.")

functions_m1 = [
    ("ajouter_producteur()", "Permet d'ajouter un nouveau producteur avec validation du nom"),
    ("afficher_producteurs()", "Affiche la liste complète avec nombre de récoltes par producteur"),
    ("rechercher_producteur(nom)", "Trouve un producteur par son nom (recherche insensible à la casse)")
]

for func, desc in functions_m1:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(f"{func}: ")
    run.bold = True
    p.add_run(desc)

add_heading(doc, "3.2 Module 2 : Gestion des Récoltes", level=3)
add_paragraph_with_spacing(doc,
    "Ce module gère l'enregistrement et le suivi des récoltes. Il implémente un système de validation "
    "stricte pour assurer la qualité des données et la cohérence du système.")

functions_m2 = [
    ("produit_valide(produit)", "Valide qu'un produit fait partie de la liste autorisée"),
    ("enregistrer_recolte()", "Enregistre une nouvelle récolte après validation complète"),
    ("afficher_recoltes(nom=None)", "Affiche toutes les récoltes ou celles d'un producteur spécifique")
]

for func, desc in functions_m2:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(f"{func}: ")
    run.bold = True
    p.add_run(desc)

add_paragraph_with_spacing(doc, "Produits autorisés : Mil, Maïs, Riz, Arachide")

add_heading(doc, "3.3 Module 3 : Statistiques de Production", level=3)
add_paragraph_with_spacing(doc,
    "Ce module fournit une vue globale de la production de la coopérative. Il agrège les données "
    "de tous les producteurs pour donner une perspective d'ensemble.")

functions_m3 = [
    ("production_totale()", "Calcule la quantité totale produite en kg par tous les producteurs"),
    ("production_par_produit()", "Décompose la production par type de produit")
]

for func, desc in functions_m3:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(f"{func}: ")
    run.bold = True
    p.add_run(desc)

add_heading(doc, "3.4 Module 4 : Analyses Avancées", level=3)
add_paragraph_with_spacing(doc,
    "Ce module fournit des analyses détaillées des performances des producteurs pour identifier "
    "les meilleures pratiques et les opportunités d'amélioration.")

functions_m4 = [
    ("meilleur_producteur(liste)", "Identifie le producteur avec la production totale la plus élevée"),
    ("moyenne_producteur(liste)", "Calcule la production moyenne par producteur")
]

for func, desc in functions_m4:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(f"{func}: ")
    run.bold = True
    p.add_run(desc)

add_heading(doc, "3.5 Module 5 : Interface Principale", level=3)
add_paragraph_with_spacing(doc,
    "Ce module fournit l'interface utilisateur du système et intègre tous les autres modules. "
    "Il gère l'interaction avec l'utilisateur via un menu interactif.")

add_paragraph_with_spacing(doc, "Options du menu :")
menu_options = [
    "1. Ajouter un producteur",
    "2. Afficher les producteurs",
    "3. Rechercher un producteur",
    "4. Enregistrer une récolte",
    "5. Afficher les récoltes",
    "6. Production totale",
    "7. Production par produit",
    "8. Meilleur producteur",
    "9. Production moyenne",
    "0. Quitter"
]

for option in menu_options:
    doc.add_paragraph(option, style='List Number')

doc.add_page_break()

# ========================
# III. ALGORITHMES
# ========================
add_heading(doc, "III. Algorithmes et Traitements Réalisés", level=1)

add_heading(doc, "1. Algorithme d'Enregistrement d'une Récolte", level=2)
add_paragraph_with_spacing(doc,
    "Cet algorithme illustre le processus complet de validation et d'enregistrement d'une récolte :")

pseudocode_recolte = """
FONCTION enregistrer_recolte()
    ENTREE : Aucune (l'utilisateur fournit les données)
    SORTIE : Récolte enregistrée ou message d'erreur
    
    DEBUT
        1. Demander le nom du producteur à l'utilisateur
        2. Rechercher le producteur dans la liste
        3. SI producteur n'existe pas ALORS
            Afficher "Le producteur n'existe pas"
            RETOURNER
        4. Demander le produit récolté
        5. SI produit non valide ALORS
            Afficher "Produit invalide"
            RETOURNER
        6. Demander la quantité en kg
        7. ESSAYER
            Convertir la quantité en nombre décimal
        EXCEPTER
            Afficher "La quantité doit être un nombre"
            RETOURNER
        8. SI quantité <= 0 ALORS
            Afficher "La quantité doit être positive"
            RETOURNER
        9. Créer structure recolte avec :
            - nom du producteur
            - nom du produit (capitalisé)
            - quantité en kg
        10. Ajouter la récolte à la liste des récoltes du producteur
        11. Afficher "Récolte enregistrée avec succès"
    FIN
"""

p = doc.add_paragraph(pseudocode_recolte)
p.style = 'List Bullet'
p.paragraph_format.left_indent = Inches(0.3)

add_heading(doc, "2. Algorithme de Calcul de Production Totale", level=2)
pseudocode_production = """
FONCTION production_totale()
    ENTREE : producteurs (liste globale)
    SORTIE : total (nombre décimal)
    
    DEBUT
        1. Initialiser total = 0
        2. POUR chaque producteur DANS producteurs FAIRE
            3. POUR chaque récolte DANS producteur.recoltes FAIRE
                4. Ajouter récolte.quantite au total
        5. Afficher "Production totale : " + total + " kg"
        6. RETOURNER total
    FIN
"""

p = doc.add_paragraph(pseudocode_production)
p.style = 'List Bullet'
p.paragraph_format.left_indent = Inches(0.3)

add_heading(doc, "3. Algorithme d'Identification du Meilleur Producteur", level=2)
pseudocode_meilleur = """
FONCTION meilleur_producteur(liste_producteurs)
    ENTREE : liste_producteurs (list)
    SORTIE : nom et production du meilleur producteur
    
    DEBUT
        1. SI liste vide ALORS
            Afficher "Aucun producteur enregistré"
            RETOURNER
        2. Initialiser max_production = -1, nom_meilleur = ""
        3. POUR chaque producteur DANS liste_producteurs FAIRE
            4. Calculer total_producteur = somme des quantités du producteur
            5. SI total_producteur > max_production ALORS
                6. Mettre à jour max_production et nom_meilleur
        7. Afficher "Meilleur producteur : " + nom_meilleur
        8. Afficher "Production : " + max_production + " kg"
    FIN
"""

p = doc.add_paragraph(pseudocode_meilleur)
p.style = 'List Bullet'
p.paragraph_format.left_indent = Inches(0.3)

add_heading(doc, "4. Flux d'Exécution Principal", level=2)
add_paragraph_with_spacing(doc, "Le programme suit le flux suivant :")

flux = """
1. Démarrage : python main.py
2. Importation de tous les modules
3. Affichage du menu principal
4. BOUCLE INFINIE
    a. Afficher les options disponibles
    b. Demander à l'utilisateur de choisir une option
    c. Selon le choix :
        - Options 1-9 : Exécuter la fonction correspondante
        - Option 0 : Quitter la boucle et arrêter le programme
    d. Retourner à l'étape 4.a
5. Fin du programme
"""

p = doc.add_paragraph(flux)
p.style = 'List Bullet'

doc.add_page_break()

# ========================
# IV. CAPTURES D'ÉCRAN / DÉMONSTRATION
# ========================
add_heading(doc, "IV. Démonstration et Captures d'Écran", level=1)

add_heading(doc, "1. Menu Principal", level=2)
add_paragraph_with_spacing(doc,
    "À son lancement, le programme affiche un menu interactif avec 10 options numérotées permettant "
    "à l'utilisateur de naviguer entre les différentes fonctionnalités.")

menu_demo = """
==========================================
 GESTION D'UNE COOPÉRATIVE AGRICOLE
==========================================
1. Ajouter un producteur
2. Afficher les producteurs
3. Rechercher un producteur
4. Enregistrer une récolte
5. Afficher les récoltes
6. Production totale
7. Production par produit
8. Meilleur producteur
9. Production moyenne par producteur
0. Quitter

Votre choix :
"""

p = doc.add_paragraph(menu_demo)
p.style = 'Normal'
p.paragraph_format.left_indent = Inches(0.5)
for run in p.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

add_heading(doc, "2. Exemple d'Utilisation Complet", level=2)
add_paragraph_with_spacing(doc,
    "Scénario : Un gestionnaire de coopérative souhaite enregistrer et analyser les récoltes de ses producteurs.")

steps = [
    ("Étape 1", "Lancer le programme et ajouter 3 producteurs : Amadou, Fatou, et Moussa"),
    ("Étape 2", "Pour chaque producteur, enregistrer plusieurs récoltes de produits différents"),
    ("Étape 3", "Afficher la liste complète des producteurs avec nombre de récoltes"),
    ("Étape 4", "Afficher toutes les récoltes ou celles d'un producteur spécifique"),
    ("Étape 5", "Voir la production totale en kg"),
    ("Étape 6", "Voir la production détaillée par produit (Mil, Maïs, Riz, Arachide)"),
    ("Étape 7", "Identifier le producteur le plus performant"),
    ("Étape 8", "Calculer la production moyenne par producteur")
]

for step, desc in steps:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(f"{step}: ")
    run.bold = True
    p.add_run(desc)

add_heading(doc, "3. Fonctionnalités Clés Démontrées", level=2)

features = [
    "Validation des données : Les noms en double sont rejetés, les produits invalides sont refusés",
    "Gestion des erreurs : Conversion de quantité, vérification de positivité",
    "Recherche efficace : Localisation rapide de producteurs par nom",
    "Rapports détaillés : Production par producteur, par produit, ou globale",
    "Interface conviviale : Menu clair et messages informatifs"
]

for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_page_break()

# ========================
# V. CONCLUSION
# ========================
add_heading(doc, "V. Conclusion", level=1)

add_heading(doc, "1. Résumé des Réalisations", level=2)
add_paragraph_with_spacing(doc,
    "Le projet AGRI-SEN a atteint tous ses objectifs. Le système est complet, fonctionnel et prêt pour "
    "utilisation. Il offre une solution complète pour la gestion des coopératives agricoles avec une "
    "architecture modulaire bien pensée.")

achievements = [
    "5 modules indépendants mais intégrés, chacun avec une responsabilité claire",
    "Interface utilisateur intuitive avec menu numéroté simple",
    "Système de validation complète garantissant la qualité des données",
    "Rapports et analyses détaillées permettant une meilleure prise de décision",
    "Code bien commenté et facilement maintenable"
]

for achievement in achievements:
    doc.add_paragraph(achievement, style='List Bullet')

add_heading(doc, "2. Difficultés Rencontrées", level=2)

difficulties = [
    ("Coordination d'équipe", "La synchronisation entre 5 développeurs a nécessité une bonne communication "
     "et une organisation claire de la répartition des tâches."),
    ("Imports circulaires", "Au début, certains modules tentaient d'importer les uns les autres, ce qui "
     "causait des erreurs. Cela a été résolu en clarifiant les dépendances."),
    ("Validation des données", "Implémenter une validation correcte à chaque étape a nécessité de tester "
     "plusieurs scénarios différents pour éviter les failles."),
    ("Structure des données", "Déterminer la meilleure structure pour stocker les producteurs et récoltes "
     "a pris du temps et plusieurs itérations.")
]

for title, description in difficulties:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(f"{title}: ")
    run.bold = True
    p.add_run(description)

add_heading(doc, "3. Perspectives d'Amélioration", level=2)

improvements = [
    ("Persistance des données", "Ajouter une base de données (SQLite, MySQL, PostgreSQL) pour que les données "
     "survivent après l'arrêt du programme. Actuellement, les données sont perdues au redémarrage."),
    ("Gestion des utilisateurs", "Implémenter un système d'authentification pour différencier les rôles "
     "(gestionnaire, producteur, auditeur) avec permissions appropriées."),
    ("Édition et suppression", "Permettre la modification ou la suppression de producteurs et récoltes, "
     "avec historique des changements pour l'audit."),
    ("Export des rapports", "Générer des rapports en PDF ou Excel pour une meilleure présentation et partage. "
     "Créer des graphiques pour visualiser les tendances de production."),
    ("Application mobile", "Développer une application mobile pour que les producteurs puissent enregistrer "
     "directement leurs récoltes sur le terrain."),
    ("Système de notification", "Ajouter des alertes pour les récoltes attendues, les producteurs inactifs, "
     "ou les anomalies détectées."),
    ("Intégration météo", "Intégrer les données météorologiques pour analyser l'impact du climat sur la production."),
    ("Interface graphique", "Remplacer le menu texte par une interface graphique avec tkinter ou PyQt "
     "pour meilleure expérience utilisateur.")
]

for title, description in improvements:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(f"{title}: ")
    run.bold = True
    p.add_run(description)

add_heading(doc, "4. Evaluation Finale", level=2)
add_paragraph_with_spacing(doc,
    "AGRI-SEN est un projet réussi qui démontre une bonne compréhension des principes de développement "
    "logiciel, notamment la modularité, la séparation des responsabilités, et la validation des données. "
    "Le code est bien structuré, facilement extensible, et prêt pour une déploiement en environnement réel. "
    "Ce projet offre une base solide sur laquelle des améliorations futures peuvent être construites.")

add_paragraph_with_spacing(doc,
    "La collaboration d'équipe a été efficace, avec chaque membre apportant sa contribution spécifique. "
    "Les défis rencontrés ont été surmontés grâce à une communication claire et une approche méthodique "
    "du problème. Le résultat final est un système complète et fonctionnel qui répondra aux besoins des "
    "coopératives agricoles.")

# Sauvegarder le document
output_path = "AGRI-SEN_Rapport_Complet.docx"
doc.save(output_path)
print(f"Rapport créé avec succès : {output_path}")
